import streamlit as st
import openai
import os
from dotenv import load_dotenv
from docx import Document
from io import BytesIO

# Load API key
load_dotenv()
api_key = os.getenv("OPENAI_API_KEY")
client = openai.OpenAI(api_key=api_key)

# Streamlit page config
st.set_page_config(page_title="AI 정리툴", layout="wide")
st.title("📑 AI 내용 정리 및 기획서 생성기")
st.markdown("영어/한글 혼합 텍스트도 자동 감지하여, **항상 한국어로 구조적으로 정리**해드립니다.")

# 세션 상태 초기화
if "report_text" not in st.session_state:
    st.session_state["report_text"] = ""
if "file_name" not in st.session_state:
    st.session_state["file_name"] = "정리결과.docx"

# Layout
col1, col2 = st.columns(2)

with col1:
    st.header("📝 입력 영역")

    # 체크박스로 목적 선택
    st.markdown("⛳ **정리 목적을 선택하세요** (기본: 내용 정리 리포트)")
    use_report = st.checkbox("🧠 자유 형식 내용 정리 리포트", value=True)
    use_idea = st.checkbox("📄 아이디어 정리용 기획서")

    # 둘 다 체크한 경우 경고
    if use_report and use_idea:
        st.warning("두 항목이 모두 선택되었습니다. '내용 정리 리포트'를 우선 적용합니다.")

    user_input = st.text_area(
        "내용을 붙여넣으세요 (예: 이메일, 회의록, 아이디어 등 - 한글/영어 혼합 가능)",
        height=400,
        placeholder="예: 지난주 바이어랑 discussed delivery delay, 내부 일정 조율도 필요함 등..."
    )

    generate_button = st.button("📄 정리하기")

with col2:
    st.header("📋 정리 결과")

    if generate_button:
        if not user_input.strip():
            st.warning("내용을 입력해주세요.")
        else:
            with st.spinner("AI가 한국어로 내용을 정리하는 중입니다..."):
                # 프롬프트 생성
                if use_report:
                    prompt = f"""
                    아래는 이메일, 회의, 대화 기록 등 다양한 내용이 섞인 자유형 텍스트입니다.
                    전체 흐름과 맥락을 분석하여, 너무 축약하지 않고 **구조적이고 명확하게 한국어로 정리**해줘.
                    핵심 논의, 결정 사항, 남은 과제, 인물 등을 자동 판단해서 ### 제목으로 항목 구분해서 써줘.

                    내용:
                    \"\"\"{user_input}\"\"\"
                    """
                    st.session_state["file_name"] = "내용정리리포트.docx"

                elif use_idea:
                    prompt = f"""
                    아래 아이디어 내용을 바탕으로 1장 분량의 기획서를 작성해줘.
                    실무자가 바로 이해할 수 있게 구성은 자유롭게 판단하되, 소제목은 ### 형식으로 구분하고,
                    내용은 깔끔하고 명확한 한국어로 정리해줘.

                    내용:
                    \"\"\"{user_input}\"\"\"
                    """
                    st.session_state["file_name"] = "기획서.docx"

                try:
                    response = client.chat.completions.create(
                        model="gpt-4o",
                        messages=[{"role": "user", "content": prompt}],
                        temperature=0.7,
                    )
                    result_text = response.choices[0].message.content.strip()
                    st.session_state["report_text"] = result_text

                except Exception as e:
                    st.error(f"에러 발생: {e}")

    # 출력 및 워드 다운로드
    if st.session_state["report_text"]:
        st.markdown(st.session_state["report_text"], unsafe_allow_html=True)

        plain_text = st.session_state["report_text"].replace("### ", "").replace("**", "")
        doc = Document()
        for line in plain_text.split("\n"):
            if line.strip() == "":
                continue
            if line.startswith("#"):
                doc.add_heading(line.replace("#", "").strip(), level=2)
            else:
                doc.add_paragraph(line.strip())

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        st.download_button(
            label="📥 워드 파일로 다운로드",
            data=buffer,
            file_name=st.session_state["file_name"],
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
